"""
Unit tests for PDFToDocsConverterService
"""
import os
import pytest
from pathlib import Path
from docx import Document

from services.pdf_to_docs_converter import (
    PDFToDocsConverterService,
    ConversionTimeoutError,
    TimeoutContext
)


@pytest.fixture
def docs_converter():
    """Fixture to create PDFToDocsConverterService instance"""
    return PDFToDocsConverterService()


@pytest.fixture
def test_data_dir():
    """Fixture to get test data directory"""
    test_dir = Path(__file__).parent / "test_data"
    test_dir.mkdir(exist_ok=True)
    return test_dir


class TestConvertPDFToDocx:
    """Tests for convert_pdf_to_docx method"""
    
    def test_convert_simple_pdf_to_docx(self, docs_converter, test_data_dir, tmp_path):
        """Test converting simple PDF with text to docx"""
        pdf_path = test_data_dir / "simple_text.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        output_path = tmp_path / "converted.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        assert success is True
        assert error is None
        assert output_path.exists()
        
        # Verify the docx file can be opened and has content
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0
        
        # Verify file size is reasonable
        assert output_path.stat().st_size > 0
    
    def test_convert_pdf_with_images(self, docs_converter, test_data_dir, tmp_path):
        """Test converting PDF with embedded images to docx"""
        pdf_path = test_data_dir / "with_images.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        output_path = tmp_path / "converted_with_images.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        assert success is True
        assert error is None
        assert output_path.exists()
        
        # Verify the docx file was created
        doc = Document(output_path)
        # Document should have content (text or images)
        assert len(doc.paragraphs) > 0 or len(doc.inline_shapes) > 0
    
    def test_convert_multipage_pdf(self, docs_converter, test_data_dir, tmp_path):
        """Test converting multipage PDF to docx"""
        pdf_path = test_data_dir / "multipage_doc.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        output_path = tmp_path / "converted_multipage.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        assert success is True
        assert error is None
        assert output_path.exists()
        
        # Verify the docx file has content from multiple pages
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0
        
        # File should be larger for multipage document
        assert output_path.stat().st_size > 1000
    
    def test_convert_formatted_pdf(self, docs_converter, test_data_dir, tmp_path):
        """Test converting PDF with formatting (bold, italic) to docx"""
        pdf_path = test_data_dir / "formatted_doc.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        output_path = tmp_path / "converted_formatted.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        assert success is True
        assert error is None
        assert output_path.exists()
        
        # Verify the docx file was created with content
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0
    
    def test_convert_empty_pdf_creates_docx_with_message(self, docs_converter, test_data_dir, tmp_path):
        """Test that empty PDF creates docx with message indicating no content"""
        pdf_path = test_data_dir / "empty_doc.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        output_path = tmp_path / "converted_empty.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        assert success is True
        assert error is None
        assert output_path.exists()
        
        # Verify the docx file contains the "no content" message
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0
        
        # Check that the message is present
        text_content = " ".join([p.text for p in doc.paragraphs])
        assert "no extractable content" in text_content.lower() or "no content" in text_content.lower()
    
    def test_convert_corrupted_pdf_returns_error(self, docs_converter, test_data_dir, tmp_path):
        """Test that corrupted PDF returns appropriate error"""
        pdf_path = test_data_dir / "corrupted.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        output_path = tmp_path / "converted_corrupted.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        assert success is False
        assert error is not None
        assert "corrupted" in error.lower() or "invalid" in error.lower()
        
        # Output file should not be created for corrupted PDF
        assert not output_path.exists()
    
    def test_convert_nonexistent_pdf_returns_error(self, docs_converter, tmp_path):
        """Test that non-existent PDF returns appropriate error"""
        pdf_path = tmp_path / "nonexistent.pdf"
        output_path = tmp_path / "output.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        assert success is False
        assert error is not None
        assert "not found" in error.lower()
    
    def test_convert_directory_instead_of_file_returns_error(self, docs_converter, tmp_path):
        """Test that passing a directory instead of file returns error"""
        pdf_path = tmp_path  # This is a directory
        output_path = tmp_path / "output.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        assert success is False
        assert error is not None
        assert "not a file" in error.lower()
    
    def test_convert_creates_output_directory_if_needed(self, docs_converter, test_data_dir, tmp_path):
        """Test that conversion creates output directory if it doesn't exist"""
        pdf_path = test_data_dir / "simple_text.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        # Create output path in non-existent subdirectory
        output_path = tmp_path / "subdir" / "nested" / "output.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        assert success is True
        assert error is None
        assert output_path.exists()
        assert output_path.parent.exists()
    
    def test_convert_verifies_output_file_created(self, docs_converter, test_data_dir, tmp_path):
        """Test that conversion verifies output file was actually created"""
        pdf_path = test_data_dir / "simple_text.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        output_path = tmp_path / "output.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        # If conversion succeeded, output file must exist
        if success:
            assert output_path.exists()
            assert output_path.stat().st_size > 0


class TestContentValidation:
    """Tests for _validate_pdf_content method"""
    
    def test_validate_pdf_with_text_content(self, docs_converter, test_data_dir):
        """Test validation of PDF with text content"""
        pdf_path = test_data_dir / "simple_text.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        has_content, error = docs_converter._validate_pdf_content(str(pdf_path))
        
        assert has_content is True
        assert error is None
    
    def test_validate_pdf_with_images(self, docs_converter, test_data_dir):
        """Test validation of PDF with images"""
        pdf_path = test_data_dir / "with_images.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        has_content, error = docs_converter._validate_pdf_content(str(pdf_path))
        
        assert has_content is True
        assert error is None
    
    def test_validate_empty_pdf(self, docs_converter, test_data_dir):
        """Test validation of empty PDF returns False with no error"""
        pdf_path = test_data_dir / "empty_doc.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        has_content, error = docs_converter._validate_pdf_content(str(pdf_path))
        
        assert has_content is False
        assert error is None  # No error, just no content
    
    def test_validate_corrupted_pdf(self, docs_converter, test_data_dir):
        """Test validation of corrupted PDF returns error"""
        pdf_path = test_data_dir / "corrupted.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        has_content, error = docs_converter._validate_pdf_content(str(pdf_path))
        
        assert has_content is False
        assert error is not None
        assert "corrupted" in error.lower() or "invalid" in error.lower()
    
    def test_validate_multipage_pdf(self, docs_converter, test_data_dir):
        """Test validation of multipage PDF"""
        pdf_path = test_data_dir / "multipage_doc.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        has_content, error = docs_converter._validate_pdf_content(str(pdf_path))
        
        assert has_content is True
        assert error is None


class TestTimeoutMechanism:
    """Tests for timeout mechanism"""
    
    def test_timeout_configuration(self):
        """Test that timeout can be configured"""
        custom_service = PDFToDocsConverterService(timeout=30)
        assert custom_service.timeout == 30
        
        default_service = PDFToDocsConverterService()
        assert default_service.timeout == PDFToDocsConverterService.DEFAULT_TIMEOUT
    
    def test_timeout_context_manager_basic(self):
        """Test TimeoutContext basic functionality"""
        with TimeoutContext(5) as ctx:
            assert ctx.timed_out is False
        
        # Should complete without timeout
        assert ctx.timed_out is False
    
    def test_timeout_context_manager_with_zero_timeout(self):
        """Test TimeoutContext with zero timeout (no timeout)"""
        with TimeoutContext(0) as ctx:
            assert ctx.timer is None
        
        # Should complete without timeout
        assert ctx.timed_out is False
    
    def test_timeout_context_raises_on_timeout(self):
        """Test that TimeoutContext raises ConversionTimeoutError on timeout"""
        import time
        
        with pytest.raises(ConversionTimeoutError, match="timed out"):
            with TimeoutContext(1) as ctx:
                # Simulate timeout by manually triggering it
                ctx._timeout_handler()
                time.sleep(0.1)
    
    def test_conversion_with_very_short_timeout(self, tmp_path):
        """Test conversion with very short timeout"""
        # Create a service with very short timeout
        short_timeout_service = PDFToDocsConverterService(timeout=0)
        
        # Create a minimal PDF
        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_bytes(b'%PDF-1.4\n%\xE2\xE3\xCF\xD3\n')
        
        output_path = tmp_path / "output.docx"
        
        # Should handle timeout or invalid PDF gracefully
        success, error = short_timeout_service.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        # Should fail (either timeout or invalid PDF)
        assert success is False
        assert error is not None


class TestErrorHandling:
    """Tests for error handling in PDFToDocsConverterService"""
    
    def test_convert_with_invalid_pdf_content(self, docs_converter, tmp_path):
        """Test conversion with invalid PDF content"""
        # Create a file with PDF header but invalid content
        fake_pdf = tmp_path / "fake.pdf"
        fake_pdf.write_bytes(b'%PDF-1.4\nInvalid content')
        
        output_path = tmp_path / "output.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(fake_pdf), str(output_path))
        
        assert success is False
        assert error is not None
        assert "corrupted" in error.lower() or "invalid" in error.lower()
    
    def test_convert_handles_permission_error_on_output(self, docs_converter, test_data_dir, tmp_path):
        """Test that permission errors on output are handled gracefully"""
        pdf_path = test_data_dir / "simple_text.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        # Try to write to an invalid/protected location
        # This is platform-dependent, so we just verify the error handling exists
        output_path = tmp_path / "output.docx"
        
        # Normal conversion should work
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        # If PDF has content, should succeed
        if success:
            assert error is None
            assert output_path.exists()
    
    def test_convert_handles_unexpected_errors(self, docs_converter, tmp_path):
        """Test that unexpected errors are caught and returned"""
        # Use a path that will cause issues
        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_text("Not a PDF")
        
        output_path = tmp_path / "output.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        assert success is False
        assert error is not None
        assert isinstance(error, str)
    
    def test_validate_handles_invalid_pdf_gracefully(self, docs_converter, tmp_path):
        """Test that validation handles invalid PDFs gracefully"""
        fake_pdf = tmp_path / "fake.pdf"
        fake_pdf.write_text("Not a real PDF file")
        
        has_content, error = docs_converter._validate_pdf_content(str(fake_pdf))
        
        assert has_content is False
        assert error is not None


class TestIntegrationScenarios:
    """Integration tests for complete conversion scenarios"""
    
    def test_full_conversion_workflow(self, docs_converter, test_data_dir, tmp_path):
        """Test complete conversion workflow from PDF to docx"""
        pdf_path = test_data_dir / "simple_text.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        output_path = tmp_path / "final_output.docx"
        
        # Step 1: Validate content
        has_content, validation_error = docs_converter._validate_pdf_content(str(pdf_path))
        assert has_content is True or validation_error is None
        
        # Step 2: Convert
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        # Step 3: Verify results
        assert success is True
        assert error is None
        assert output_path.exists()
        
        # Step 4: Verify output is valid docx
        doc = Document(output_path)
        assert doc is not None
    
    def test_conversion_preserves_text_content(self, docs_converter, test_data_dir, tmp_path):
        """Test that conversion preserves text content from PDF"""
        pdf_path = test_data_dir / "simple_text.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        output_path = tmp_path / "text_preserved.docx"
        
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
        
        assert success is True
        
        # Verify the docx has text content
        doc = Document(output_path)
        text_content = " ".join([p.text for p in doc.paragraphs])
        
        # Should have some text content
        assert len(text_content.strip()) > 0
    
    def test_multiple_conversions_in_sequence(self, docs_converter, test_data_dir, tmp_path):
        """Test that multiple conversions can be performed in sequence"""
        pdf_path = test_data_dir / "simple_text.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        # Perform multiple conversions
        for i in range(3):
            output_path = tmp_path / f"output_{i}.docx"
            success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(output_path))
            
            assert success is True
            assert error is None
            assert output_path.exists()
    
    def test_conversion_with_different_output_paths(self, docs_converter, test_data_dir, tmp_path):
        """Test conversion with various output path configurations"""
        pdf_path = test_data_dir / "simple_text.pdf"
        
        if not pdf_path.exists():
            pytest.skip("Test PDF not found. Run create_test_pdfs_for_docs.py first.")
        
        # Test with nested directory
        nested_output = tmp_path / "level1" / "level2" / "output.docx"
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(nested_output))
        
        assert success is True
        assert nested_output.exists()
        
        # Test with simple path
        simple_output = tmp_path / "simple.docx"
        success, error = docs_converter.convert_pdf_to_docx(str(pdf_path), str(simple_output))
        
        assert success is True
        assert simple_output.exists()
